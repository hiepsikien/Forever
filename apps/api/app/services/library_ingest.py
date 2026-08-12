"""Gemini classify + OCR for library document ingest."""

from __future__ import annotations

import base64
import json
import logging
import mimetypes
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import httpx
from nanoid import generate
from sqlalchemy.orm import Session

from ..config import get_settings
from ..models import LibraryIngestJob, LibraryIngestProposal
from ..services.poetry_clean import enrich_poem, format_body, format_body_tts, clean_body_lines
from ..services.storage import absolute_media_path

logger = logging.getLogger(__name__)

ALLOWED_THEMES = {
    "vo_chong",
    "con_cai",
    "gia_dinh",
    "nghe_giao",
    "tho",
    "biet_on",
    "truyen_thong",
}
ALLOWED_KINDS = {"poem", "milestone", "note", "knowledge"}

SYSTEM = """\
Bạn là trợ lý Forever — số hóa tài liệu gia đình thành đề nghị đưa vào Thư viện.
Hard rules:
- Chỉ đọc nội dung có trong tài liệu. KHÔNG bịa thơ, mốc đời, hay sự kiện.
- Tách thành các món riêng (mỗi bài thơ / mỗi mốc / mỗi ghi chú).
- kind chỉ được: poem | milestone | note | knowledge.
- Thơ: mỗi câu một dòng; meter: luc_bat|song_that_luc_bat|that_ngon|other|unknown.
- Album thơ tặng: tên người làm thơ / họa thường đứng DƯỚI bài (có khi kèm nơi, năm).
  BẮT BUỘC điền field "attribution" đúng như trên trang (vd. "Vũ Kiểm — Hoàng Mai, Hà Nội").
  Không cắt, không bỏ, không đưa attribution vào title trừ khi đó là tên bài.
- authorship: "gift" nếu bài do người khác viết tặng/họa/mừng (có attribution hoặc
  lời đề tặng); "own" nếu là thơ của chính người được nhớ / không có dấu hiệu tặng.
- Theme whitelist: vo_chong, con_cai, gia_dinh, nghe_giao, tho, biet_on, truyen_thong.
- occurred_at dạng YYYY-MM-DD nếu biết năm/ngày; không chắc thì null.
- Tối đa 8 món mỗi lần trả lời — ưu tiên bài thơ/mốc rõ ràng; bỏ phần mơ hồ.
- Trả ĐÚNG một JSON object ngắn gọn, không markdown fence.
"""

USER = """\
Đọc tài liệu và đề nghị các món đưa vào Thư viện. Trả JSON (tối đa 8 items).
Với thơ tặng/họa: luôn điền attribution + authorship=\"gift\". Thơ của chính người được nhớ: authorship=\"own\".

{
  "items": [
    {
      "kind": "poem|milestone|note|knowledge",
      "title": "…",
      "body": "…",
      "attribution": "Tên người làm thơ / Theo … / nơi, năm — hoặc \\"\\"",
      "authorship": "own|gift",
      "meter": "luc_bat|…|unknown",
      "themes": ["gia_dinh"],
      "occurred_at": "YYYY-MM-DD hoặc null"
    }
  ],
  "notes": ""
}
"""


def _salvage_items_from_truncated_json(blob: str) -> dict | None:
    """Pull complete item objects out of a truncated Gemini JSON array."""
    marker = blob.find('"items"')
    if marker < 0:
        return None
    bracket = blob.find("[", marker)
    if bracket < 0:
        return None
    decoder = json.JSONDecoder()
    items: list[Any] = []
    idx = bracket + 1
    n = len(blob)
    while idx < n:
        while idx < n and blob[idx] in " \t\r\n,":
            idx += 1
        if idx >= n or blob[idx] == "]":
            break
        if blob[idx] != "{":
            break
        try:
            obj, end = decoder.raw_decode(blob, idx)
        except json.JSONDecodeError:
            break
        if isinstance(obj, dict):
            items.append(obj)
        idx = end
    if not items:
        return None
    return {"items": items, "notes": "salvaged_truncated_json"}


def _parse_json_object(text: str) -> dict:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        lines = [ln for ln in cleaned.splitlines() if not ln.strip().startswith("```")]
        cleaned = "\n".join(lines).strip()
    start, end = cleaned.find("{"), cleaned.rfind("}")
    if start < 0:
        raise RuntimeError(f"No JSON object in model output: {cleaned[:300]}")
    # Prefer full object when the closing brace exists; else salvage from truncation.
    blob = cleaned[start : (end + 1 if end > start else len(cleaned))]
    try:
        return json.loads(blob)
    except json.JSONDecodeError:
        salvaged = _salvage_items_from_truncated_json(cleaned[start:])
        if salvaged:
            logger.warning(
                "library ingest salvaged %s items from truncated JSON",
                len(salvaged["items"]),
            )
            return salvaged
        raise RuntimeError(
            f"JSON hỏng từ model (bị cắt giữa chừng). Đoạn đầu: {cleaned[start:start + 200]}"
        )


def _extract_gemini_payload(data: dict) -> tuple[str, str]:
    """Return (text, finish_reason)."""
    candidates = data.get("candidates") or []
    if not candidates:
        raise RuntimeError(f"No candidates: {json.dumps(data)[:400]}")
    cand = candidates[0]
    finish = str(cand.get("finishReason") or cand.get("finish_reason") or "")
    parts = ((cand.get("content") or {}).get("parts")) or []
    texts = [p.get("text", "") for p in parts if isinstance(p, dict) and p.get("text")]
    text = "\n".join(texts).strip()
    if not text:
        raise RuntimeError(f"Empty Gemini text (finish={finish or '?'})")
    return text, finish


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_MIME = "application/msword"
MAX_WORD_CHARS = 120_000
# Keep chunks small so Gemini can finish JSON for dense poetry albums.
WORD_CHUNK_CHARS = 3_500
OLE_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"

_DOC_JUNK_MARKERS = (
    "INCLUDEPICTURE",
    "MERGEFORMAT",
    "ADOBE PHOTOSHOP",
    "DEFAULT PARAGRAPH",
    "WORDDOCUMENT",
    "SUMMARYINFORMATION",
    "MSODATASTORE",
    "TIMES NEW ROMAN",
    "CAMBRIA",
    "NORMAL.DOT",
    "BALLOON TEXT",
    "DOCUMENT MAP",
    "TABLE NORMAL",
    "HEADING 1",
    "KẾT QUẢ HÌNH ẢNH",
    "IMAGE RESULT",
    "HTTP://",
    "HTTPS://",
    ".JPG",
    ".PNG",
    ".GIF",
    ".JPEG",
    ".VNCOMMERCIAL",
)


def _clip_word_text(text: str, *, empty_hint: str) -> str:
    text = (text or "").strip()
    if not text:
        raise RuntimeError(empty_hint)
    if len(text) > MAX_WORD_CHARS:
        text = text[:MAX_WORD_CHARS] + "\n…[đã cắt bớt vì quá dài]"
    return text


def _viet_char_ratio(text: str) -> float:
    if not text:
        return 0.0
    viet = sum(
        1
        for c in text
        if ("\u1ea0" <= c <= "\u1ef9") or (c.lower() in "ăâêôơưđ")
    )
    return viet / max(len(text), 1)


def _text_quality_score(text: str) -> float:
    """Higher is better. textutil often returns huge binary-looking UTF-16 junk."""
    text = (text or "").strip()
    if not text:
        return 0.0
    sample = text[:8000]
    printable = sum(1 for c in sample if c.isprintable() or c in "\n\t")
    letters = sum(1 for c in sample if c.isalpha())
    nulish = sample.count("\x00") + sum(1 for c in sample if ord(c) < 9)
    score = (printable / max(len(sample), 1)) * 0.45
    score += min(letters / 400.0, 1.0) * 0.25
    score += min(_viet_char_ratio(sample) * 8.0, 1.0) * 0.35
    score -= min(nulish / 50.0, 1.0) * 0.5
    # Penalize enormous dumps that are mostly non-letters (OLE mis-decode).
    if len(text) > 200_000 and letters < 2000:
        score *= 0.1
    return max(0.0, score)


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Thiếu python-docx trên server — không đọc được file DOCX."
        ) from exc
    doc = Document(str(path))
    chunks: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            chunks.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                chunks.append(line)
    return _clip_word_text(
        "\n".join(chunks),
        empty_hint=(
            "DOCX không có chữ để đọc (có thể chỉ chứa ảnh). "
            "Xuất PDF hoặc chụp từng trang rồi tải lên."
        ),
    )


def _decode_tool_text(raw: bytes | str) -> str:
    if isinstance(raw, str):
        data = raw.encode("latin-1", errors="replace")
    else:
        data = raw
    if data.startswith(b"\xff\xfe") or data.startswith(b"\xfe\xff"):
        return data.decode("utf-16", errors="replace")
    # UTF-16LE without BOM is common from macOS textutil on Vietnamese Word.
    if b"\x00" in data[:200]:
        try:
            return data.decode("utf-16-le", errors="replace")
        except UnicodeDecodeError:
            pass
    for enc in ("utf-8", "utf-16-le", "cp1258", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _is_usable_doc_line(line: str) -> bool:
    cleaned = " ".join(line.split()).strip()
    if len(cleaned) < 2 or len(cleaned) > 320:
        return False
    upper = cleaned.upper()
    if any(marker in upper for marker in _DOC_JUNK_MARKERS):
        return False
    if upper.startswith("{\\*"):
        return False
    if sum(1 for c in cleaned if ord(c) < 32) > 2:
        return False
    letters = sum(1 for c in cleaned if c.isalpha())
    if letters < 2:
        return False
    uniq = set(cleaned)
    if len(cleaned) > 10 and len(uniq) <= 3:
        return False
    if len(cleaned) > 12:
        dominant = max(cleaned.count(c) for c in uniq)
        if dominant / len(cleaned) > 0.55:
            return False
    # Filename / URL slug without spaces.
    if " " not in cleaned and "-" in cleaned and len(cleaned) >= 10:
        alnum = sum(1 for c in cleaned if c.isalnum() or c in "-_./")
        if alnum / len(cleaned) > 0.9:
            return False
    weird = sum(
        1
        for c in cleaned
        if ord(c) > 0x024F
        and not (0x1EA0 <= ord(c) <= 0x1EF9)
        and not (0x2010 <= ord(c) <= 0x2027)
    )
    if weird >= 2:
        return False
    return True


def _clean_extracted_text(text: str) -> str:
    text = text.replace("\x00", "")
    lines: list[str] = []
    prev = None
    for line in text.splitlines():
        cleaned = " ".join(line.split()).strip()
        if not cleaned or cleaned == prev:
            continue
        if not _is_usable_doc_line(cleaned):
            continue
        lines.append(cleaned)
        prev = cleaned
    # Drop trailing metadata blobs (fonts / image alt text) after poetry.
    while len(lines) > 40:
        tail = "\n".join(lines[-40:])
        head = "\n".join(lines[:-40])
        if _viet_char_ratio(head) >= 0.04 and _viet_char_ratio(tail) < 0.015:
            lines = lines[:-20]
            continue
        break
    return "\n".join(lines).strip()


def _extract_doc_ole_utf16(path: Path) -> str:
    """Scan OLE .doc binary for UTF-16LE runs — works when textutil dumps garbage."""
    data = path.read_bytes()
    if not data.startswith(OLE_MAGIC):
        return ""
    chunks: list[str] = []
    i = 0
    n = len(data)
    while i + 1 < n:
        chars: list[str] = []
        j = i
        while j + 1 < n:
            lo, hi = data[j], data[j + 1]
            cp = lo | (hi << 8)
            if hi == 0 and (cp in (9, 10, 13) or 32 <= cp < 127):
                chars.append("\n" if cp in (10, 13) else ("\t" if cp == 9 else chr(cp)))
                j += 2
                continue
            if (
                0x00C0 <= cp <= 0x024F
                or 0x1EA0 <= cp <= 0x1EF9
                or 0x0100 <= cp <= 0x017F
                or cp in (0x2013, 0x2014, 0x2018, 0x2019, 0x201C, 0x201D, 0x2026)
            ):
                chars.append(chr(cp))
                j += 2
                continue
            break
        if len(chars) >= 12:
            piece = "".join(chars)
            letters = sum(1 for c in piece if c.isalpha() or c.isspace())
            if letters / max(len(piece), 1) >= 0.7 and sum(c.isalpha() for c in piece) >= 8:
                chunks.append(piece)
            i = j
        else:
            i += 1
    return _clean_extracted_text("\n".join(chunks))


def _run_cmd_bytes(argv: list[str], *, timeout: float = 60.0) -> tuple[int, bytes, str]:
    import subprocess

    try:
        proc = subprocess.run(
            argv,
            capture_output=True,
            timeout=timeout,
            check=False,
        )
    except FileNotFoundError:
        return 127, b"", "not found"
    except subprocess.TimeoutExpired:
        return 124, b"", "timeout"
    err = (proc.stderr or b"").decode("utf-8", errors="replace")
    return proc.returncode, proc.stdout or b"", err


def _extract_doc_text(path: Path) -> str:
    """Legacy .doc — OS tools first, then OLE UTF-16 scan for image-heavy albums."""
    empty_hint = (
        "DOC không có chữ để đọc (có thể chỉ chứa ảnh). "
        "Xuất DOCX/PDF hoặc chụp từng trang rồi tải lên."
    )
    tried: list[str] = []
    best_text = ""
    best_score = 0.0

    def consider(label: str, text: str) -> None:
        nonlocal best_text, best_score
        cleaned = _clean_extracted_text(text)
        score = _text_quality_score(cleaned)
        tried.append(f"{label}(chars={len(cleaned)},q={score:.2f})")
        if score > best_score and cleaned:
            best_text, best_score = cleaned, score

    for argv in (
        ["textutil", "-convert", "txt", "-encoding", "UTF-8", "-stdout", str(path)],
        ["textutil", "-convert", "txt", "-stdout", str(path)],
    ):
        code, out_b, err = _run_cmd_bytes(argv, timeout=180.0)
        if code == 127:
            break
        if code == 0 and out_b.strip():
            consider("textutil", _decode_tool_text(out_b))
        else:
            tried.append(f"textutil({code}:{(err or '')[:80]})")

    for argv in (
        ["antiword", "-w", "0", str(path)],
        ["catdoc", "-w", str(path)],
    ):
        code, out_b, err = _run_cmd_bytes(argv, timeout=180.0)
        if code == 127:
            continue
        if code == 0 and out_b.strip():
            consider(argv[0], _decode_tool_text(out_b))
        else:
            tried.append(f"{argv[0]}({code}:{(err or '')[:80]})")

    import tempfile

    with tempfile.TemporaryDirectory(prefix="forever-doc-") as tmp:
        for bin_name in ("soffice", "libreoffice"):
            code, out_b, err = _run_cmd_bytes(
                [
                    bin_name,
                    "--headless",
                    "--norestore",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    tmp,
                    str(path),
                ],
                timeout=180.0,
            )
            if code == 127:
                continue
            txt_path = Path(tmp) / f"{path.stem}.txt"
            if txt_path.is_file():
                consider(bin_name, _decode_tool_text(txt_path.read_bytes()))
            else:
                tried.append(f"{bin_name}({code}:{(err or '')[:80]})")

    # Image-heavy Word albums often make textutil emit ~100MB binary; OLE scan wins.
    if best_score < 0.45 or path.read_bytes()[:8] == OLE_MAGIC:
        ole_text = _extract_doc_ole_utf16(path)
        if ole_text:
            consider("ole-utf16", ole_text)

    if best_text and best_score >= 0.25:
        return _clip_word_text(best_text, empty_hint=empty_hint)

    detail = "; ".join(tried) if tried else "không có textutil/antiword/catdoc/LibreOffice"
    raise RuntimeError(
        "Không đọc được file .doc trên máy chủ. "
        f"({detail}). Cài antiword hoặc LibreOffice, hoặc xuất DOCX/PDF."
    )


def _split_word_chunks(text: str, *, size: int = WORD_CHUNK_CHARS) -> list[str]:
    text = (text or "").strip()
    if not text:
        return []
    if len(text) <= size:
        return [text]
    lines = text.splitlines()
    chunks: list[str] = []
    buf: list[str] = []
    buf_len = 0
    for line in lines:
        add = len(line) + 1
        if buf and buf_len + add > size:
            chunks.append("\n".join(buf))
            buf, buf_len = [], 0
        buf.append(line)
        buf_len += add
    if buf:
        chunks.append("\n".join(buf))
    return chunks


def _gemini_classify(*, path: Path, mime: str, text: str | None = None) -> dict:
    settings = get_settings()
    api_key = (settings.gemini_api_key or "").strip()
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY chưa cấu hình.")
    model = settings.gemini_model
    api_base = settings.gemini_api_base.rstrip("/")
    url = f"{api_base}/models/{model}:generateContent"

    def call_once(parts: list[dict], *, max_tokens: int = 8192) -> dict:
        payload = {
            "systemInstruction": {"parts": [{"text": SYSTEM}]},
            "contents": [{"role": "user", "parts": parts}],
            "generationConfig": {
                "temperature": 0.1,
                "maxOutputTokens": max_tokens,
                "thinkingConfig": {"thinkingBudget": 0},
                "responseMimeType": "application/json",
            },
        }
        with httpx.Client(timeout=240.0) as client:
            res = client.post(
                url,
                params={"key": api_key},
                headers={"Content-Type": "application/json"},
                json=payload,
            )
            res.raise_for_status()
            body = res.json()
            text_out, finish = _extract_gemini_payload(body)
            parsed = _parse_json_object(text_out)
            if finish.upper() == "MAX_TOKENS":
                logger.warning(
                    "library ingest Gemini hit MAX_TOKENS; kept %s items",
                    len(parsed.get("items") or []),
                )
            return parsed

    def classify_text_chunk(chunk: str, *, part_label: str) -> dict:
        parts = [
            {"text": USER},
            {
                "text": (
                    f"{part_label}"
                    "Nội dung đã trích từ tài liệu Word (chỉ chữ, không có ảnh nhúng):\n\n"
                    f"{chunk}"
                )
            },
        ]
        try:
            return call_once(parts)
        except RuntimeError as exc:
            # One more try at half size if JSON still blows up.
            if len(chunk) < 800 or "JSON" not in str(exc):
                raise
            logger.warning("library ingest chunk failed (%s); splitting", exc)
            mid = len(chunk) // 2
            split_at = chunk.rfind("\n", 0, mid) or mid
            left = chunk[:split_at].strip()
            right = chunk[split_at:].strip()
            merged: list[Any] = []
            notes: list[str] = []
            for piece in (left, right):
                if not piece:
                    continue
                raw = call_once(
                    [
                        {"text": USER},
                        {
                            "text": (
                                f"{part_label}"
                                "Nội dung đã trích từ tài liệu Word "
                                "(chỉ chữ, không có ảnh nhúng):\n\n"
                                f"{piece}"
                            )
                        },
                    ]
                )
                merged.extend(raw.get("items") or [])
                note = (raw.get("notes") or "").strip()
                if note:
                    notes.append(note)
            return {"items": merged, "notes": " | ".join(notes)}

    if text is None:
        raw = path.read_bytes()
        if len(raw) > 100 * 1024 * 1024:
            raise RuntimeError("File quá lớn để xử lý (tối đa ~100 MB).")
        b64 = base64.b64encode(raw).decode("ascii")
        return call_once(
            [{"text": USER}, {"inline_data": {"mime_type": mime, "data": b64}}]
        )

    chunks = _split_word_chunks(text)
    if not chunks:
        raise RuntimeError("Không có chữ để phân loại.")
    merged_items: list[Any] = []
    notes: list[str] = []
    total = len(chunks)
    for idx, chunk in enumerate(chunks, start=1):
        prefix = ""
        if total > 1:
            prefix = (
                f"Đây là phần {idx}/{total} của cùng một tài liệu. "
                "Chỉ đề nghị món từ phần này (tối đa 8); bỏ qua phần đã xử lý trước.\n\n"
            )
        raw = classify_text_chunk(chunk, part_label=prefix)
        merged_items.extend(raw.get("items") or [])
        note = (raw.get("notes") or "").strip()
        if note:
            notes.append(note)
    return {"items": merged_items, "notes": " | ".join(notes)}

def _parse_occurred(value: Any) -> datetime | None:
    if not value or not isinstance(value, str):
        return None
    raw = value.strip()
    if not raw:
        return None
    try:
        if len(raw) == 4 and raw.isdigit():
            return datetime(int(raw), 1, 1, tzinfo=timezone.utc)
        parsed = datetime.fromisoformat(raw.replace("Z", "+00:00"))
        if parsed.tzinfo is None:
            parsed = parsed.replace(tzinfo=timezone.utc)
        return parsed
    except ValueError:
        return None


def _attribution_text(item: dict) -> str:
    raw = (
        item.get("attribution")
        or item.get("author")
        or item.get("byline")
        or item.get("nguoi_lam_tho")
        or ""
    )
    if isinstance(raw, list):
        parts = [str(x).strip() for x in raw if str(x).strip()]
        return "\n".join(parts)
    return str(raw or "").strip()


def _append_attribution(body: str, attribution: str) -> str:
    attr = "\n".join(ln.strip() for ln in attribution.splitlines() if ln.strip())
    if not attr:
        return body
    compact_body = " ".join(body.lower().split())
    compact_attr = " ".join(attr.lower().split())
    if compact_attr and compact_attr in compact_body:
        return body
    return f"{body.rstrip()}\n\n{attr}".strip()


def _normalize_items(raw: dict) -> list[dict]:
    out: list[dict] = []
    for item in raw.get("items") or []:
        kind = (item.get("kind") or "note").strip().lower()
        if kind not in ALLOWED_KINDS:
            kind = "note"
        title = (item.get("title") or "").strip()[:200]
        body = (item.get("body") or "").strip()
        attribution = _attribution_text(item)
        if not body:
            continue
        themes = [t for t in (item.get("themes") or []) if t in ALLOWED_THEMES]
        meter = (item.get("meter") or "unknown").strip() or "unknown"
        authorship_raw = (item.get("authorship") or "").strip().lower()
        if authorship_raw in {"gift", "tang", "gifted", "tặng"}:
            authorship = "gift"
        elif authorship_raw in {"own", "cua_minh", "self"}:
            authorship = "own"
        elif attribution:
            authorship = "gift"
        else:
            authorship = (
                "gift" if _looks_like_gift(title, body, attribution) else "own"
            )
        row: dict[str, Any] = {
            "kind": kind,
            "title": title,
            "body": body,
            "meter": meter,
            "themes": themes,
            "authorship": authorship if kind == "poem" else "own",
            "occurred_at": _parse_occurred(item.get("occurred_at")),
            "body_tts": "",
        }
        if kind == "poem":
            enriched = enrich_poem(
                {
                    "title": title,
                    "body": body,
                    "meter": meter,
                    "themes": themes,
                }
            )
            lines = clean_body_lines(enriched.get("body") or body, meter=meter)
            row["body"] = _append_attribution(format_body(lines), attribution)
            row["body_tts"] = format_body_tts(lines, meter=meter)
            row["title"] = (enriched.get("title") or title)[:200]
            row["themes"] = enriched.get("themes") or themes
            row["meter"] = enriched.get("meter") or meter
            composed = enriched.get("composed_on")
            if composed and not row["occurred_at"]:
                row["occurred_at"] = _parse_occurred(composed)
        else:
            row["body"] = _append_attribution(body, attribution)
        out.append(row)
    return out


def _looks_like_gift(title: str, body: str, attribution: str) -> bool:
    hay = f"{title}\n{body}\n{attribution}".lower()
    return any(
        tip in hay
        for tip in (
            "tặng",
            "mừng",
            "họa",
            "theo ",
            "kính tặng",
            "thơ mừng",
            "chúc mừng",
        )
    )


def process_library_ingest_job(db: Session, job_id: str) -> None:
    job = db.query(LibraryIngestJob).filter(LibraryIngestJob.id == job_id).one_or_none()
    if not job:
        return
    if job.status not in {"queued", "running"}:
        return
    now = datetime.now(timezone.utc)
    job.status = "running"
    job.started_at = now
    job.error_message = ""
    db.commit()

    try:
        path = Path(absolute_media_path(job.input_path))
        if not path.is_file():
            raise RuntimeError("Không tìm thấy file đã tải lên.")
        mime = job.input_mime or mimetypes.guess_type(str(path))[0] or "application/octet-stream"
        suffix = path.suffix.lower()
        if suffix == ".docx":
            mime = DOCX_MIME
        elif suffix == ".doc":
            mime = DOC_MIME
        if mime.startswith("image/") or mime == "application/pdf":
            raw = _gemini_classify(path=path, mime=mime)
        elif mime == DOCX_MIME or suffix == ".docx":
            word_text = _extract_docx_text(path)
            raw = _gemini_classify(path=path, mime=DOCX_MIME, text=word_text)
        elif mime == DOC_MIME or suffix == ".doc":
            word_text = _extract_doc_text(path)
            raw = _gemini_classify(path=path, mime=DOC_MIME, text=word_text)
        else:
            raise RuntimeError(
                "Định dạng chưa hỗ trợ. Dùng ảnh (.jpg/.png), PDF, DOCX hoặc DOC."
            )
        items = _normalize_items(raw)
        # Replace any prior proposals if re-run.
        for old in list(job.proposals):
            db.delete(old)
        db.flush()
        for i, item in enumerate(items):
            prop = LibraryIngestProposal(
                id=generate(),
                job_id=job.id,
                kind=item["kind"],
                title=item["title"],
                body=item["body"],
                body_tts=item.get("body_tts") or "",
                meter=item.get("meter") or "",
                themes_json=json.dumps(item.get("themes") or [], ensure_ascii=False),
                authorship=item.get("authorship") or "own",
                occurred_at=item.get("occurred_at"),
                identity_id=job.identity_id,
                review_status="pending",
                memory_item_id=None,
                sort_order=i,
                created_at=datetime.now(timezone.utc),
            )
            db.add(prop)
        job.status = "needs_review"
        job.finished_at = datetime.now(timezone.utc)
        job.model = get_settings().gemini_model
        db.commit()
    except Exception as exc:
        logger.exception("library ingest failed for %s", job_id)
        job.status = "failed"
        job.error_message = str(exc)[:800]
        job.finished_at = datetime.now(timezone.utc)
        db.commit()
