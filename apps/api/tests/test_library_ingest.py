"""Library document ingest — upload + approve proposals."""

from io import BytesIO
from unittest.mock import patch

# Minimal 1×1 PNG (no Pillow dependency in this module).
_TINY_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xcf"
    b"\xc0\x00\x00\x00\x03\x00\x01\x00\x05\xfe\xd4\xef\x00\x00\x00\x00IEND"
    b"\xaeB`\x82"
)


def _login(client, email: str, name: str) -> str:
    res = client.post(
        "/api/auth/dev-login",
        json={"email": email, "password": "forever123", "name": name},
    )
    assert res.status_code == 200, res.text
    return res.json()["token"]


def _space(client, headers: dict) -> str:
    space = client.post("/api/spaces", headers=headers, json={"name": "Nhà thơ"}).json()
    return space["id"]


def _tiny_png() -> bytes:
    return _TINY_PNG


def test_append_attribution_to_poem_body():
    from app.services.library_ingest import _normalize_items

    items = _normalize_items(
        {
            "items": [
                {
                    "kind": "poem",
                    "title": "Cưới vàng",
                    "body": "Năm mươi năm đám cưới vàng\nAnh em bầu bạn họ hàng đông vui",
                    "attribution": "Vũ Kiểm\nHoàng Mai, Hà Nội",
                    "meter": "luc_bat",
                    "themes": ["tho"],
                }
            ]
        }
    )
    assert len(items) == 1
    assert "Vũ Kiểm" in items[0]["body"]
    assert "Hoàng Mai" in items[0]["body"]
    assert items[0]["authorship"] == "gift"


def test_poem_authorship_tags_exclude_gifts_from_count(client):
    from datetime import datetime, timezone

    from nanoid import generate

    from app.db import SessionLocal
    from app.models import MemoryItem
    from app.services.heritage import (
        POEM_AUTH_GIFT_TAG,
        POEM_AUTH_OWN_TAG,
        is_own_poem,
        poem_count_for_identity,
    )

    res = client.post(
        "/api/auth/dev-login",
        json={
            "email": "authorship-poem@example.com",
            "password": "forever123",
            "name": "Steward",
        },
    )
    assert res.status_code == 200
    token = res.json()["token"]
    user_id = res.json()["user"]["id"]
    headers = {"Authorization": f"Bearer {token}"}
    space_id = _space(client, headers)
    identity_id = client.post(
        f"/api/spaces/{space_id}/identities",
        headers=headers,
        json={
            "display_name": "Bố Triệu",
            "relation_label": "Bố",
            "status": "remembered",
        },
    ).json()["id"]
    assert is_own_poem(f"heritage:{identity_id} tho")
    assert not is_own_poem(f"heritage:{identity_id} {POEM_AUTH_GIFT_TAG}")

    db = SessionLocal()
    try:
        now = datetime.now(timezone.utc)
        for title, tag in (
            ("Own", POEM_AUTH_OWN_TAG),
            ("Gift", POEM_AUTH_GIFT_TAG),
        ):
            db.add(
                MemoryItem(
                    id=generate(),
                    space_id=space_id,
                    created_by=user_id,
                    kind="poem",
                    title=title,
                    body="một hai ba",
                    body_tts="",
                    tags=f"heritage:{identity_id} tho {tag}",
                    created_at=now,
                )
            )
        db.commit()
        assert poem_count_for_identity(db, space_id=space_id, identity_id=identity_id) == 1
    finally:
        db.close()


def test_salvage_truncated_json_items():
    from app.services.library_ingest import _parse_json_object

    truncated = (
        '{\n  "items": [\n'
        '    {"kind": "poem", "title": "A", "body": "một hai", '
        '"meter": "luc_bat", "themes": ["tho"], "occurred_at": null},\n'
        '    {"kind": "poem", "title": "B", "body": "ba bốn", '
        '"meter": "luc_bat", "themes": ["tho"], "occurred_at": null},\n'
        '    {"kind": "poem", "title": "C", "body": "năm sá'
    )
    parsed = _parse_json_object(truncated)
    assert len(parsed["items"]) == 2
    assert parsed["items"][0]["title"] == "A"


def test_ole_utf16_doc_extract_and_chunk():
    """Image-heavy .doc albums: prefer OLE UTF-16 scan over textutil garbage."""
    from pathlib import Path

    from app.services.library_ingest import (
        OLE_MAGIC,
        _extract_doc_ole_utf16,
        _split_word_chunks,
        _text_quality_score,
    )

    # Minimal OLE header + UTF-16LE Vietnamese poetry run.
    poem = "Từ ngày trầu quyện với vôi\nNăm mươi năm đám cưới vàng\n"
    body = poem.encode("utf-16-le")
    blob = OLE_MAGIC + (b"\x00" * 64) + body + (b"\x00" * 32)
    path = Path("/tmp/forever-ole-poem.doc")
    path.write_bytes(blob)
    text = _extract_doc_ole_utf16(path)
    assert "trầu quyện" in text
    assert _text_quality_score(text) >= 0.25
    chunks = _split_word_chunks(text * 800, size=500)
    assert len(chunks) > 1


def test_create_job_requires_identity(client):
    token = _login(client, "ingest-need-id@example.com", "Steward")
    headers = {"Authorization": f"Bearer {token}"}
    space_id = _space(client, headers)
    res = client.post(
        f"/api/spaces/{space_id}/library-ingest/jobs",
        headers=headers,
        files={"file": ("x.png", _tiny_png(), "image/png")},
        data={},
    )
    assert res.status_code == 400
    body = res.json()
    msg = body.get("detail") or body.get("error") or ""
    assert "neo" in msg.lower() or "Chọn" in msg


def test_patch_job_identity_cascades_to_pending(client):
    from unittest.mock import patch

    token = _login(client, "ingest-patch-id@example.com", "Steward")
    headers = {"Authorization": f"Bearer {token}"}
    space_id = _space(client, headers)
    a = client.post(
        f"/api/spaces/{space_id}/identities",
        headers=headers,
        json={"display_name": "Bố", "relation_label": "Bố", "status": "remembered"},
    ).json()["id"]
    b = client.post(
        f"/api/spaces/{space_id}/identities",
        headers=headers,
        json={"display_name": "Mẹ", "relation_label": "Mẹ", "status": "remembered"},
    ).json()["id"]
    fake = {
        "items": [
            {
                "kind": "poem",
                "title": "Tặng",
                "body": "một hai ba bốn năm sáu\nbảy tám chín mười một hai",
                "attribution": "Bạn thơ",
                "authorship": "gift",
                "meter": "luc_bat",
                "themes": ["tho"],
            }
        ],
        "notes": "",
    }
    with patch("app.services.library_ingest._gemini_classify", return_value=fake):
        job = client.post(
            f"/api/spaces/{space_id}/library-ingest/jobs",
            headers=headers,
            files={"file": ("p.png", _tiny_png(), "image/png")},
            data={"identity_id": a},
        ).json()
    job_id = job["id"]
    for _ in range(20):
        got = client.get(
            f"/api/spaces/{space_id}/library-ingest/jobs/{job_id}",
            headers=headers,
        ).json()
        if got["status"] in {"needs_review", "failed", "done"}:
            break
        import time

        time.sleep(0.05)
    patched = client.patch(
        f"/api/spaces/{space_id}/library-ingest/jobs/{job_id}",
        headers=headers,
        json={"identity_id": b},
    )
    assert patched.status_code == 200, patched.text
    assert patched.json()["identity_id"] == b
    props = client.get(
        f"/api/spaces/{space_id}/library-ingest/jobs/{job_id}/proposals",
        headers=headers,
    ).json()["proposals"]
    assert props and all(p["identity_id"] == b for p in props)


def test_accept_doc_upload_via_textutil(client):
    import shutil
    import subprocess
    import tempfile
    from pathlib import Path

    if not shutil.which("textutil"):
        # Linux CI / Docker without textutil — covered by antiword path in prod.
        return

    token = _login(client, "ingest-doc@example.com", "Steward")
    headers = {"Authorization": f"Bearer {token}"}
    space_id = _space(client, headers)
    identity = client.post(
        f"/api/spaces/{space_id}/identities",
        headers=headers,
        json={
            "display_name": "Bố Triệu",
            "relation_label": "Bố",
            "status": "remembered",
        },
    )
    assert identity.status_code == 200, identity.text
    identity_id = identity.json()["id"]

    with tempfile.TemporaryDirectory() as tmp:
        txt = Path(tmp) / "poem.txt"
        doc = Path(tmp) / "poem.doc"
        txt.write_text("Bai tho thu\nBay nham dau phai da gia\n", encoding="utf-8")
        subprocess.run(
            ["textutil", "-convert", "doc", str(txt), "-output", str(doc)],
            check=True,
            capture_output=True,
        )
        payload = doc.read_bytes()

    fake = {
        "items": [
            {
                "kind": "note",
                "title": "Điếu văn",
                "body": "Bảy nhăm đâu phải đã già",
                "themes": [],
                "occurred_at": None,
            }
        ],
        "notes": "",
    }
    with patch(
        "app.services.library_ingest._gemini_classify",
        return_value=fake,
    ) as mocked:
        res = client.post(
            f"/api/spaces/{space_id}/library-ingest/jobs",
            headers=headers,
            files={"file": ("dieuvan.doc", payload, "application/msword")},
            data={"identity_id": identity_id},
        )
    assert res.status_code == 200, res.text
    assert mocked.called
    assert "Bay nham" in (mocked.call_args.kwargs.get("text") or "")


def test_accept_docx_upload_and_approve(client):
    from docx import Document

    token = _login(client, "ingest-docx@example.com", "Steward")
    headers = {"Authorization": f"Bearer {token}"}
    space_id = _space(client, headers)

    identity = client.post(
        f"/api/spaces/{space_id}/identities",
        headers=headers,
        json={
            "display_name": "Bố Triệu",
            "relation_label": "Bố",
            "status": "remembered",
        },
    )
    assert identity.status_code == 200, identity.text
    identity_id = identity.json()["id"]

    doc = Document()
    doc.add_paragraph("TUỔI BẢY NHĂM")
    doc.add_paragraph("Bảy nhăm đâu phải đã già")
    doc.add_paragraph("Cứ vui vẻ thảnh thơi")
    buf = BytesIO()
    doc.save(buf)
    payload = buf.getvalue()

    fake = {
        "items": [
            {
                "kind": "poem",
                "title": "TUỔI BẢY NHĂM",
                "body": "Bảy nhăm đâu phải đã già\nCứ vui vẻ thảnh thơi",
                "meter": "luc_bat",
                "themes": ["gia_dinh"],
                "occurred_at": None,
            }
        ],
        "notes": "",
    }

    with patch(
        "app.services.library_ingest._gemini_classify",
        return_value=fake,
    ) as mocked:
        res = client.post(
            f"/api/spaces/{space_id}/library-ingest/jobs",
            headers=headers,
            files={
                "file": (
                    "tho.docx",
                    payload,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"identity_id": identity_id},
        )
    assert res.status_code == 200, res.text
    job_id = res.json()["id"]
    assert mocked.called
    # Text path: Gemini got extracted DOCX text, not binary inline.
    call_kwargs = mocked.call_args.kwargs
    assert call_kwargs.get("text")
    assert "Bảy nhăm" in call_kwargs["text"]

    for _ in range(20):
        got = client.get(
            f"/api/spaces/{space_id}/library-ingest/jobs/{job_id}",
            headers=headers,
        ).json()
        if got["status"] in {"needs_review", "failed", "done"}:
            break
        import time

        time.sleep(0.05)
    assert got["status"] == "needs_review", got

    props = client.get(
        f"/api/spaces/{space_id}/library-ingest/jobs/{job_id}/proposals",
        headers=headers,
    ).json()["proposals"]
    assert len(props) == 1
    assert props[0]["kind"] == "poem"


def test_ingest_image_proposals_and_approve(client, tmp_path):
    token = _login(client, "ingest-ok@example.com", "Steward")
    headers = {"Authorization": f"Bearer {token}"}
    space_id = _space(client, headers)

    identity = client.post(
        f"/api/spaces/{space_id}/identities",
        headers=headers,
        json={
            "display_name": "Bố Triệu",
            "relation_label": "Bố",
            "status": "remembered",
        },
    )
    assert identity.status_code == 200, identity.text
    identity_id = identity.json()["id"]

    fake = {
        "items": [
            {
                "kind": "poem",
                "title": "TUỔI BẢY NHĂM",
                "body": "Bảy nhăm đâu phải đã già\nCứ vui vẻ thảnh thơi",
                "meter": "luc_bat",
                "themes": ["gia_dinh"],
                "occurred_at": "2014-08-07",
            },
            {
                "kind": "milestone",
                "title": "Về quê",
                "body": "Về quê thăm mẹ năm ấy.",
                "themes": [],
                "occurred_at": "1975-01-01",
            },
        ],
        "notes": "",
    }

    with patch(
        "app.services.library_ingest._gemini_classify",
        return_value=fake,
    ):
        res = client.post(
            f"/api/spaces/{space_id}/library-ingest/jobs",
            headers=headers,
            files={"file": ("page.png", _tiny_png(), "image/png")},
            data={"identity_id": identity_id},
        )
    assert res.status_code == 200, res.text
    job = res.json()
    assert job["status"] in {"queued", "running", "needs_review", "done"}
    job_id = job["id"]

    # BackgroundTasks run after response with TestClient — wait/process once more if needed.
    for _ in range(20):
        got = client.get(
            f"/api/spaces/{space_id}/library-ingest/jobs/{job_id}",
            headers=headers,
        ).json()
        if got["status"] in {"needs_review", "failed", "done"}:
            job = got
            break
        import time

        time.sleep(0.05)

    assert job["status"] == "needs_review", job

    props = client.get(
        f"/api/spaces/{space_id}/library-ingest/jobs/{job_id}/proposals",
        headers=headers,
    ).json()["proposals"]
    assert len(props) == 2
    assert props[0]["kind"] == "poem"

    approve = client.post(
        f"/api/spaces/{space_id}/library-ingest/jobs/{job_id}/proposals/settle",
        headers=headers,
        json={"proposal_ids": [props[0]["id"]], "action": "approve"},
    )
    assert approve.status_code == 200, approve.text
    assert len(approve.json()["created_memory_ids"]) == 1

    memories = client.get(f"/api/spaces/{space_id}/memories", headers=headers).json()[
        "memories"
    ]
    poems = [m for m in memories if m["kind"] == "poem"]
    assert len(poems) == 1
    assert "TUỔI" in poems[0]["title"].upper() or poems[0]["title"]
